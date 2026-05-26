#!/usr/bin/env python3
"""Build final DOCX/TXT files from confirmed Markdown drafts."""

from __future__ import annotations

import argparse
import html
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from common import ensure_dir, read_json, safe_filename

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


BLACK_RGB = "000000"


def strip_markdown_links(text: str) -> str:
    text = re.sub(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"<(https?://[^>]+)>", r"\1", text)
    return text


def parse_application_lines(md_path: Path) -> tuple[list[str], list[str]]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    fields = [line.strip() for line in lines if line.strip().startswith("➤")]
    warnings = [line for line in fields if "待用户确认" in line]
    return fields, warnings


def parse_application_field(md_path: Path, field_name: str) -> str:
    if not md_path.exists():
        return ""
    prefix = f"➤{field_name}："
    for line in md_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith(prefix):
            return stripped[len(prefix) :].strip()
    return ""


def application_version(draft_dir: Path) -> str:
    version = parse_application_field(draft_dir / "申请表信息.md", "版本号")
    if "待用户确认" in version:
        return ""
    return version


def application_software_name(draft_dir: Path) -> str:
    name = parse_application_field(draft_dir / "申请表信息.md", "软件全称")
    if "待用户确认" in name:
        return ""
    return name


def write_application_txt(draft_dir: Path, out_dir: Path) -> tuple[Path | None, list[str]]:
    md_path = draft_dir / "申请表信息.md"
    if not md_path.exists():
        return None, ["缺少草稿/申请表信息.md"]
    fields, warnings = parse_application_lines(md_path)
    out_path = out_dir / "申请表信息.txt"
    out_path.write_text("\n".join(fields) + "\n", encoding="utf-8")
    return out_path, warnings


def read_json_if_exists(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    return read_json(path)


def confirmation_issues(workdir: Path) -> list[str]:
    draft_dir = workdir / "草稿"
    issues: list[str] = []

    business = read_json_if_exists(draft_dir / "业务理解.json")
    if not business or not business.get("user_confirmed"):
        issues.append("业务理解尚未确认：请确认 草稿/业务理解.md 后记录 `business` 门禁")

    selection = read_json_if_exists(draft_dir / "代码文件选择.json")
    if not selection or not selection.get("user_confirmed"):
        issues.append("代码文件选择尚未确认：请确认 草稿/代码文件选择.json 后记录 `code-selection` 门禁")

    screenshot = read_json_if_exists(workdir / "截图方式确认.json")
    if not screenshot.get("screenshot_method_confirmed"):
        issues.append("截图方式尚未确认：请选择截图方式后记录 `screenshot-method` 门禁")

    app_md = draft_dir / "申请表信息.md"
    if app_md.exists():
        _, warnings = parse_application_lines(app_md)
        if warnings:
            issues.append("申请表信息仍包含“待用户确认”字段")
    else:
        issues.append("缺少 草稿/申请表信息.md")

    app_confirmation = read_json_if_exists(draft_dir / "申请表字段确认.json")
    if not app_confirmation.get("application_fields_confirmed"):
        issues.append("申请表字段尚未确认：请补全字段后记录 `application-fields` 门禁")

    markdown_confirmation = read_json_if_exists(draft_dir / "最终生成确认.json")
    if not markdown_confirmation.get("markdown_confirmed"):
        issues.append("Markdown 草稿尚未最终确认：请确认全部草稿后记录 `markdown` 门禁")

    return issues


def parse_code_pages(md_path: Path) -> list[tuple[int, list[str]]]:
    pages: list[tuple[int, list[str]]] = []
    current_no: int | None = None
    current_lines: list[str] = []
    in_fence = False

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        page_match = re.match(r"^##\s+第\s*(\d+)\s*页", raw.strip())
        if page_match:
            if current_no is not None:
                pages.append((current_no, current_lines))
            current_no = int(page_match.group(1))
            current_lines = []
            in_fence = False
            continue
        if raw.strip().startswith("```"):
            in_fence = not in_fence
            continue
        if current_no is not None and in_fence:
            current_lines.append(raw)

    if current_no is not None:
        pages.append((current_no, current_lines))
    return pages


def set_run_font(run: Any, name: str, size_pt: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def set_normal_font(document: Any, name: str = "SimSun", size_pt: float = 10.5) -> None:
    style = document.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    try:
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass


def set_style_black(document: Any) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        try:
            document.styles[style_name].font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass


def force_black_document(document: Any) -> None:
    set_style_black(document)
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            except Exception:
                                pass


def configure_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.4)


def configure_code_a4(document: Any) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)


def add_page_field(paragraph: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, result, end):
        run = paragraph.add_run()
        run._r.append(element)
        set_run_font(run, "SimSun", 8)


def set_code_header(document: Any, software_name: str, version: str) -> None:
    section = document.sections[0]
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(12)
    prefix = paragraph.add_run(f"{software_name} {version}    第 ")
    set_run_font(prefix, "SimSun", 8)
    add_page_field(paragraph)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, "SimSun", 8)


def build_code_docx_python(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    pages = parse_code_pages(md_path)
    if not pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")

    document = Document()
    configure_code_a4(document)
    set_normal_font(document, "Consolas", 7.2)
    set_style_black(document)
    set_code_header(document, software_name, version)

    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(12)
            run = p.add_run(line if line else " ")
            set_run_font(run, "Consolas", 7.2)
        if index != len(pages) - 1:
            document.add_page_break()

    force_black_document(document)
    document.save(out_path)


def paragraph_xml(text: str, font: str = "SimSun", size_half_points: int = 21, align: str | None = None, line_twips: int = 240) -> str:
    align_xml = f'<w:jc w:val="{align}"/>' if align else ""
    escaped = html.escape(text)
    return (
        "<w:p>"
        f"<w:pPr>{align_xml}<w:spacing w:after=\"0\" w:line=\"{line_twips}\" w:lineRule=\"exact\"/></w:pPr>"
        "<w:r>"
        f"<w:rPr><w:rFonts w:ascii=\"{font}\" w:hAnsi=\"{font}\" w:eastAsia=\"{font}\"/>"
        f"<w:color w:val=\"{BLACK_RGB}\"/>"
        f"<w:sz w:val=\"{size_half_points}\"/><w:szCs w:val=\"{size_half_points}\"/></w:rPr>"
        f"<w:t xml:space=\"preserve\">{escaped}</w:t>"
        "</w:r>"
        "</w:p>"
    )


def page_break_xml() -> str:
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'


def page_field_runs_xml() -> str:
    return (
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:t>1</w:t></w:r>'
        '<w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/>'
        f'<w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
        '<w:fldChar w:fldCharType="end"/></w:r>'
    )


def header_xml(header_text: str) -> str:
    escaped = html.escape(header_text)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p>
    <w:pPr><w:jc w:val="right"/><w:spacing w:after="0" w:line="240" w:lineRule="exact"/></w:pPr>
    <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve">{escaped}    第 </w:t></w:r>
    {page_field_runs_xml()}
    <w:r><w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr><w:t xml:space="preserve"> 页</w:t></w:r>
  </w:p>
</w:hdr>"""


def minimal_docx(out_path: Path, body_xml: str, header_text: str | None = None) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    header_rel = (
        '<Relationship Id="rIdHeader1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/>'
        if header_text
        else ""
    )
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{header_rel}</Relationships>"""
    styles = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr><w:rFonts w:ascii="SimSun" w:hAnsi="SimSun" w:eastAsia="SimSun"/><w:color w:val="{BLACK_RGB}"/><w:sz w:val="21"/></w:rPr>
  </w:style>
</w:styles>"""
    document = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    {body_xml}
      <w:sectPr>
        {'<w:headerReference w:type="default" r:id="rIdHeader1"/>' if header_text else ''}
        <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="680" w:right="567" w:bottom="680" w:left="567" w:header="283" w:footer="283" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/styles.xml", styles)
        zf.writestr("word/document.xml", document)
        if header_text:
            zf.writestr("word/header1.xml", header_xml(header_text))


def force_black_xml(xml: str) -> str:
    xml = re.sub(r"<w:hyperlink\b[^>]*>", "", xml)
    xml = xml.replace("</w:hyperlink>", "")
    xml = re.sub(r"<w:color\b[^>]*/>", f'<w:color w:val="{BLACK_RGB}"/>', xml)

    def ensure_rpr_color(match: re.Match[str]) -> str:
        value = match.group(0)
        if "<w:color" in value:
            return value
        return value.replace("</w:rPr>", f'<w:color w:val="{BLACK_RGB}"/></w:rPr>')

    xml = re.sub(r"<w:rPr\b[^>]*>.*?</w:rPr>", ensure_rpr_color, xml, flags=re.S)
    xml = re.sub(r"<w:r>(?!<w:rPr>)", f'<w:r><w:rPr><w:color w:val="{BLACK_RGB}"/></w:rPr>', xml)
    return xml


def normalize_docx_text_color(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    color_xml_parts = (
        "word/document.xml",
        "word/styles.xml",
        "word/numbering.xml",
        "word/header",
        "word/footer",
    )
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename.endswith(".xml") and item.filename.startswith(color_xml_parts):
                text = data.decode("utf-8")
                data = force_black_xml(text).encode("utf-8")
            elif item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="ignore")
                if "hyperlink" in text:
                    text = re.sub(r'\s*<Relationship\b[^>]*Type="[^"]*/hyperlink"[^>]*/>', "", text)
                    data = text.encode("utf-8")
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def build_code_docx_ooxml(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    pages = parse_code_pages(md_path)
    if not pages:
        raise RuntimeError(f"No code pages parsed from {md_path}")
    body: list[str] = []
    for index, (page_no, lines) in enumerate(pages):
        for line in lines:
            body.append(paragraph_xml(line if line else " ", font="Consolas", size_half_points=14, line_twips=240))
        if index != len(pages) - 1:
            body.append(page_break_xml())
    minimal_docx(out_path, "\n".join(body), header_text=f"{software_name} {version}")


def add_markdown_table(document: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, text in enumerate(rows[0]):
        table.rows[0].cells[idx].text = strip_markdown_links(text)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(cells)]):
            cells[idx].text = strip_markdown_links(text)


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_image(document: Any, image_path: Path) -> None:
    if not image_path.exists():
        p = document.add_paragraph()
        run = p.add_run(f"[截图缺失：{image_path}]")
        set_run_font(run, "SimSun", 10.5)
        return
    try:
        document.add_picture(str(image_path), width=Inches(5.8))
    except Exception:
        p = document.add_paragraph()
        run = p.add_run(f"[截图无法插入：{image_path}]")
        set_run_font(run, "SimSun", 10.5)


def build_manual_docx_python(md_path: Path, out_path: Path, base_dir: Path) -> None:
    document = Document()
    configure_a4(document)
    set_normal_font(document, "SimSun", 10.5)
    set_style_black(document)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    table_buf: list[list[str]] = []
    in_fence = False

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            data = [row for row in table_buf if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
            add_markdown_table(document, data)
            table_buf = []

    for line in lines:
        stripped = line.strip()
        stripped = strip_markdown_links(stripped)
        if stripped.startswith("```"):
            flush_table()
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if stripped.startswith("<!--") and "截图" in stripped:
            stripped = "【截图预留：请在此处插入当前功能页面或操作结果截图。】"
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buf.append(parse_table_line(stripped))
            continue
        flush_table()
        if not stripped:
            continue
        image_match = re.search(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, (base_dir / image_match.group(1)).resolve())
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            p = document.add_heading(heading.group(2), level=level)
            for run in p.runs:
                try:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception:
                    pass
            continue
        if re.match(r"^[-*+]\s+", stripped):
            p = document.add_paragraph(style="List Bullet")
            run = p.add_run(re.sub(r"^[-*+]\s+", "", stripped))
            set_run_font(run, "SimSun", 10.5)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            run = p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            set_run_font(run, "SimSun", 10.5)
            continue
        p = document.add_paragraph()
        run = p.add_run(stripped)
        set_run_font(run, "SimSun", 10.5)
    flush_table()
    force_black_document(document)
    document.save(out_path)


def pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def build_with_pandoc(md_path: Path, out_path: Path, code_mode: bool = False) -> None:
    if not pandoc_available():
        raise RuntimeError("python-docx is unavailable and pandoc is not installed")
    source = md_path
    tmp_name: str | None = None
    original_text = md_path.read_text(encoding="utf-8")
    text = original_text
    text = re.sub(r"```text\s*\nSTOP_FOR_USER\n.*?```", "", text, flags=re.S)
    text = re.sub(r"<!--[^>]*截图[^>]*-->", "【截图预留：请在此处插入当前功能页面或操作结果截图。】", text)
    text = strip_markdown_links(text)
    if code_mode:
        text = re.sub(r"(?=^##\s+第\s*\d+\s*页)", r"\n\\newpage\n", text, flags=re.M)
    if code_mode or "STOP_FOR_USER" in original_text:
        with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
            tmp.write(text)
            tmp_name = tmp.name
        source = Path(tmp_name)
    try:
        subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", str(source), "-o", str(out_path)], check=True)
    finally:
        if tmp_name:
            Path(tmp_name).unlink(missing_ok=True)


def build_code_docx(md_path: Path, out_path: Path, software_name: str, version: str) -> None:
    if DOCX_AVAILABLE:
        build_code_docx_python(md_path, out_path, software_name, version)
    else:
        build_code_docx_ooxml(md_path, out_path, software_name, version)
    normalize_docx_text_color(out_path)


def build_manual_docx(md_path: Path, out_path: Path, base_dir: Path) -> None:
    if DOCX_AVAILABLE:
        build_manual_docx_python(md_path, out_path, base_dir)
    else:
        build_with_pandoc(md_path, out_path, code_mode=False)
    normalize_docx_text_color(out_path)


def run_command(command: list[str], cwd: Path | None = None, timeout: int = 60) -> tuple[int, str]:
    try:
        completed = subprocess.run(command, cwd=cwd, text=True, capture_output=True, timeout=timeout)
        return completed.returncode, (completed.stdout + completed.stderr).strip()
    except Exception as exc:
        return 99, str(exc)


def docx_checks(skill_dir: Path, outputs: list[Path]) -> list[str]:
    notes: list[str] = []
    env_script = skill_dir / "vendor/docx-toolkit/scripts/env_check.sh"
    preview_script = skill_dir / "vendor/docx-toolkit/scripts/docx_preview.sh"
    if env_script.exists():
        code, output = run_command(["bash", str(env_script)], cwd=env_script.parent.parent, timeout=30)
        status = "READY" if code == 0 else "NOT READY"
        first_lines = "\n".join(output.splitlines()[:12])
        notes.append(f"DOCX env: {status}\n\n```text\n{first_lines}\n```")
    else:
        notes.append("DOCX env: vendor script missing")

    if preview_script.exists():
        for out in outputs:
            code, output = run_command(["bash", str(preview_script), str(out)], timeout=45)
            first_lines = "\n".join(output.splitlines()[:8])
            notes.append(f"Preview {out.name}: exit={code}\n\n```text\n{first_lines}\n```")
    return notes


def build_all(workdir: Path, software_name: str, version: str, skip_preview: bool) -> dict[str, Any]:
    workdir = ensure_dir(workdir)
    draft_dir = workdir / "草稿"
    final_dir = ensure_dir(workdir / "正式资料")
    app_name = application_software_name(draft_dir)
    app_version = application_version(draft_dir)
    final_software_name = app_name or software_name
    final_version = app_version or version
    safe_name = safe_filename(final_software_name)
    outputs: list[Path] = []
    warnings: list[str] = []
    if app_name and app_name != software_name:
        warnings.append(f"命令参数软件名称为 {software_name}，正式资料已按申请表信息软件名称 {app_name} 生成")
    if app_version and app_version != version:
        warnings.append(f"命令参数版本号为 {version}，正式资料已按申请表信息版本号 {app_version} 生成")
    screenshot_confirmation = read_json_if_exists(workdir / "截图方式确认.json")
    screenshot_method = screenshot_confirmation.get("screenshot_method")
    screenshot_manifest = workdir / "截图/截图清单.json"
    if screenshot_method == "skip":
        warnings.append("用户选择暂不截图；操作手册已保留截图预留位置")
    elif screenshot_method and not screenshot_manifest.exists():
        warnings.append("操作手册截图未生成或未插入；操作手册应保留截图预留位置")
    elif screenshot_manifest.exists():
        screenshots = read_json_if_exists(screenshot_manifest).get("screenshots") or []
        if not screenshots:
            warnings.append("操作手册截图清单为空；操作手册应保留截图预留位置")

    app_txt, app_warnings = write_application_txt(draft_dir, final_dir)
    if app_txt:
        outputs.append(app_txt)
    warnings.extend(app_warnings)

    code_specs = [
        ("代码-前30页.md", f"{safe_name}-代码(前30页).docx"),
        ("代码-后30页.md", f"{safe_name}-代码(后30页).docx"),
        ("代码-全部.md", f"{safe_name}-代码(全部).docx"),
    ]
    for md_name, docx_name in code_specs:
        md_path = draft_dir / md_name
        if md_path.exists():
            out_path = final_dir / docx_name
            build_code_docx(md_path, out_path, final_software_name, final_version)
            outputs.append(out_path)

    manual_md = draft_dir / "操作手册.md"
    if manual_md.exists():
        manual_out = final_dir / f"{safe_name}_操作手册.docx"
        manual_source = manual_md
        tmp_manual: Path | None = None
        if app_name and app_name != software_name:
            text = manual_md.read_text(encoding="utf-8").replace(software_name, app_name)
            with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as tmp:
                tmp.write(text)
                tmp_manual = Path(tmp.name)
            manual_source = tmp_manual
        try:
            build_manual_docx(manual_source, manual_out, draft_dir)
        finally:
            if tmp_manual:
                tmp_manual.unlink(missing_ok=True)
        outputs.append(manual_out)
    else:
        warnings.append("缺少草稿/操作手册.md")

    skill_dir = Path(__file__).resolve().parents[1]
    notes = [] if skip_preview else docx_checks(skill_dir, [p for p in outputs if p.suffix.lower() == ".docx"])
    report = write_report(final_dir, outputs, warnings, notes)
    return {"outputs": [str(p) for p in outputs], "warnings": warnings, "report": str(report)}


def write_report(workdir: Path, outputs: list[Path], warnings: list[str], notes: list[str]) -> Path:
    report = workdir / "生成报告.md"
    lines = ["# 生成报告", "", "## 输出文件", ""]
    for path in outputs:
        size = path.stat().st_size if path.exists() else 0
        lines.append(f"- `{path.name}` ({size} bytes)")
    lines.extend(["", "## 警告", ""])
    if warnings:
        lines.extend(f"- {warning}" for warning in warnings)
    else:
        lines.append("- 无")
    lines.extend(["", "## DOCX 校验", ""])
    if notes:
        lines.extend(notes)
    else:
        lines.append("- 已跳过预览校验")
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workdir", default="软件著作权申请资料")
    parser.add_argument("--software-name", required=True)
    parser.add_argument("--version", default="V1.0")
    parser.add_argument("--skip-preview", action="store_true")
    args = parser.parse_args()

    workdir = Path(args.workdir)
    issues = confirmation_issues(workdir)
    if issues:
        print("STOP_FOR_USER")
        print("NEXT_ACTION: 正式 Word/TXT 生成前必须完成以下确认：")
        for issue in issues:
            print(f"- {issue}")
        raise SystemExit(2)

    result = build_all(workdir, args.software_name, args.version, args.skip_preview)
    print(f"OK final materials: {Path(args.workdir) / '正式资料'}")
    for output in result["outputs"]:
        print(output)
    if result["warnings"]:
        print("Warnings:")
        for warning in result["warnings"]:
            print(f"- {warning}")
    print(f"Report: {result['report']}")


if __name__ == "__main__":
    main()
